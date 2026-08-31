#!/usr/bin/env python3
"""PDF 转 DOCX 工具 - 先尝试提取文本层，失败则对每页做 OCR（支持扫描版/照片式 PDF）"""
import sys
import argparse
import fitz  # pymupdf
from docx import Document
from docx.shared import Pt

from pp_ocr_onnx import ocr_pdf as ocr_pdf_func


def extract_text_layer(pdf_path: str) -> str:
    """尝试用 pymupdf 提取文本层"""
    text_parts = []
    try:
        doc = fitz.open(pdf_path)
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
    except Exception:
        pass
    return "\n".join(text_parts)


def ocr_pdf(pdf_path: str, lang: str = "chi_sim+eng") -> str:
    """扫描版 PDF：每页渲染成图后 OCR"""
    return ocr_pdf_func(pdf_path, lang)


def text_to_docx(text: str, output_path: str) -> None:
    doc = Document()
    for line in text.splitlines():
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.size = Pt(11)
    doc.save(output_path)


def pdf_to_docx(pdf_path: str, output_path: str) -> bool:
    # 1. 先提取文本层
    text = extract_text_layer(pdf_path)
    # 2. 文本层为空 -> 扫描版/照片式 PDF，走 OCR
    if not text.strip():
        sys.stderr.write("文本层为空，启用 OCR 识别扫描版 PDF...\n")
        text = ocr_pdf(pdf_path)
    if not text.strip():
        return False
    text_to_docx(text, output_path)
    return True


def main():
    parser = argparse.ArgumentParser(description="PDF to DOCX converter")
    parser.add_argument("input", help="Input PDF file path")
    parser.add_argument("output", help="Output DOCX file path")
    args = parser.parse_args()

    success = pdf_to_docx(args.input, args.output)
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
