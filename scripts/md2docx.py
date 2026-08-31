#!/usr/bin/env python3
"""Convert markdown content to DOCX using python-docx."""
import sys
import argparse
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def markdown_to_docx(md_path: str, output_path: str) -> bool:
    try:
        with open(md_path, "r", encoding="utf-8") as f:
            lines = f.readlines()
    except Exception as e:
        sys.stderr.write(f"读取 markdown 失败: {e}\n")
        return False

    doc = Document()
    in_code_block = False
    code_lines = []
    list_stack = []  # track nested lists

    for line in lines:
        line = line.rstrip("\n\r")

        # Code blocks
        if line.startswith("```"):
            if in_code_block:
                # End code block
                for cl in code_lines:
                    p = doc.add_paragraph(cl)
                    p.paragraph_format.left_indent = Inches(0.5)
                    for run in p.runs:
                        run.font.name = "Courier New"
                        run.font.size = Pt(9)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # Headings
        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=0)
            continue
        if line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=1)
            continue
        if line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=2)
            continue
        if line.startswith("#### "):
            p = doc.add_heading(line[5:].strip(), level=3)
            continue

        # Empty line
        if not line.strip():
            doc.add_paragraph()
            continue

        # Unordered list
        if line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(line[2:].strip(), style="List Bullet")
            continue

        # Ordered list
        import re
        ol_match = re.match(r"^(\d+)\.\s+(.*)", line)
        if ol_match:
            p = doc.add_paragraph(ol_match.group(2).strip(), style="List Number")
            continue

        # Bold/italic inline (basic)
        text = line
        p = doc.add_paragraph(text)
        # Simple bold detection: **text**
        parts = []
        rest = text
        while "**" in rest:
            idx = rest.index("**")
            end = rest.index("**", idx + 2) if end := rest.find("**", idx + 2) != -1 else -1
            if end == -1:
                break
            parts.append((rest[:idx], False))
            parts.append((rest[idx+2:end], True))
            rest = rest[end+2:]
        if parts:
            p.clear()
            for content, bold in parts:
                run = p.add_run(content)
                run.bold = bold
                run.font.size = Pt(11)
            continue

        # Plain paragraph
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.size = Pt(11)

    # Handle unclosed code block
    if code_lines:
        for cl in code_lines:
            p = doc.add_paragraph(cl)
            p.paragraph_format.left_indent = Inches(0.5)
            for run in p.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(9)

    doc.save(output_path)
    return True


def main():
    parser = argparse.ArgumentParser(description="Markdown to DOCX converter")
    parser.add_argument("input", help="Input markdown file path")
    parser.add_argument("output", help="Output DOCX file path")
    args = parser.parse_args()

    success = markdown_to_docx(args.input, args.output)
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
